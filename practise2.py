import os
import streamlit as st
import whisper
from docx import Document
from collections import Counter
import re

# Set page configuration
st.set_page_config(page_title="Audio Transcriber", page_icon="🎤", layout="wide")

# Streamlit UI Styling
st.markdown("""
    <style>
        .stTextInput>div>div>input { font-size: 16px; }
        .stButton>button { width: 100%; font-size: 16px; border-radius: 10px; }
        .stMarkdown { font-size: 16px; }
        .stAlert { font-size: 16px; }
    </style>
""", unsafe_allow_html=True)

# Streamlit UI
st.title("🎙️ Audio Transcriber")
st.write("Upload an audio file and generate a transcript using Whisper.")

# File uploader
audio_file = st.file_uploader("📂 Upload an audio file", type=["mp3", "wav", "m4a"])

# Function to generate a Word file
def save_transcript_to_docx(transcript, filename="transcript.docx"):
    doc = Document()
    doc.add_heading("Audio Transcript", level=1)
    doc.add_paragraph(transcript)
    doc_path = os.path.join("temp_audio", filename)
    doc.save(doc_path)
    return doc_path

# Function to analyze word frequency
def get_word_frequency(text):
    words = re.findall(r'\b\w+\b', text.lower())
    word_counts = Counter(words)
    return word_counts.most_common(10)

# Sidebar for additional functionalities
with st.sidebar:
    st.header("⚙️ Options")
    word_freq = st.checkbox("📊 Show Word Frequency")
    st.markdown("---")

if st.button("📝 Transcribe Audio"):
    if audio_file:
        try:
            # Save uploaded file
            temp_dir = "temp_audio"
            os.makedirs(temp_dir, exist_ok=True)
            file_path = os.path.join(temp_dir, audio_file.name)
            with open(file_path, "wb") as f:
                f.write(audio_file.read())

            # Load Whisper model
            with st.spinner("📝 Transcribing audio..."):
                model = whisper.load_model("tiny")
                result = model.transcribe(file_path)
                transcript = result["text"]

            st.success("✅ Transcription completed!")
            st.subheader("📜 Transcript")
            st.text_area("", transcript, height=300)

            # Word Frequency Option
            if word_freq:
                st.subheader("📊 Most Frequent Words")
                word_counts = get_word_frequency(transcript)
                st.write(word_counts)

            # Save transcript as Word file
            docx_path = save_transcript_to_docx(transcript)

            # Download button for transcript
            with open(docx_path, "rb") as file:
                st.download_button(
                    label="📥 Download Transcript (Word)",
                    data=file,
                    file_name="transcript.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error(f"❌ Error: {str(e)}")
    else:
        st.warning("⚠️ Please upload an audio file.")

# Footer
st.markdown("---")
st.markdown("📌 **Developed by Ali Haider** | Powered by OpenAI Whisper & Streamlit")
